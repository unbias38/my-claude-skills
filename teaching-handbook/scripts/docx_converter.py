# /// script
# requires-python = ">=3.11"
# dependencies = [
#     "mammoth",
#     "beautifulsoup4",
#     "python-docx",
# ]
# ///
import mammoth
import os
import sys
import argparse
import tempfile
import html as html_lib
from docx import Document

try:
    import style_injector
    from _polish import SIDEBAR_SEARCH_SCRIPT, add_lazy_loading
except ImportError:
    sys.path.append(os.path.dirname(os.path.abspath(__file__)))
    import style_injector
    from _polish import SIDEBAR_SEARCH_SCRIPT, add_lazy_loading


def extract_colored_runs(docx_path):
    """Return [(text, hex_color)] for runs with explicit font color."""
    doc = Document(docx_path)
    found = []

    def scan(paragraphs):
        for para in paragraphs:
            for run in para.runs:
                text = run.text
                if not text.strip():
                    continue
                color = run.font.color
                if color and color.rgb:
                    found.append((text, str(color.rgb).upper()))

    scan(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                scan(cell.paragraphs)
    return found


def recover_colors(html, colored_runs):
    """Wrap colored text in <span style="color:#..."> — only if the
    (HTML-escaped) text appears exactly once, to avoid mis-matching."""
    merged = {}
    for text, color in colored_runs:
        merged.setdefault(text, color)

    for text, color in merged.items():
        needle = html_lib.escape(text)
        if html.count(needle) == 1:
            html = html.replace(
                needle,
                f'<span style="color:#{color}">{needle}</span>',
            )
        else:
            print(f"  skipped color recovery (text not unique): {text[:30]!r}")
    return html

def convert_docx_to_html(docx_path, output_filename, page_title="教學手冊", sidebar_title="教學手冊導航"):
    print(f"Converting {docx_path} to {output_filename}...")
    
    if not os.path.exists(docx_path):
        print(f"Error: Input file '{docx_path}' not found.")
        sys.exit(1)

    with open(docx_path, "rb") as docx_file:
        result = mammoth.convert_to_html(docx_file)
        html_content = result.value
        messages = result.messages

        for msg in messages:
            print(f"Mammoth msg: {msg}")

    colored = extract_colored_runs(docx_path)
    if colored:
        print(f"Recovering {len(colored)} colored text run(s) stripped by mammoth...")
        html_content = recover_colors(html_content, colored)

    # Local addition: lazy-load all images
    html_content = add_lazy_loading(html_content)

    # Wrap in specific container for our CSS to work
    # Our CSS targets .WordSection1, so let's wrap the content in that.
    # Also add standard HTML boilerplate
    
    raw_html = f"""
    <!DOCTYPE html>
    <html lang="zh-TW">
    <head>
        <meta charset="utf-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>{page_title}</title>
        <!-- Mammoth typically produces CLEAN html, so no weird Word styles to fight against -->
        <style>
            /* Basic resets to match our expected look before style_injector takes over */
            img {{ max-width: 100%; height: auto; }}
            table {{ border-collapse: collapse; width: 100%; }}
            td, th {{ border: 1px solid #ddd; padding: 8px; }}
        </style>
    </head>
    <body>
        <div class="WordSection1">
            {html_content}
        </div>
        {SIDEBAR_SEARCH_SCRIPT}
    </body>
    </html>
    """
    
    # Write temp file for style_injector to consume
    with tempfile.NamedTemporaryFile(
        mode="w", encoding="utf-8", suffix=".html", delete=False
    ) as f:
        temp_file = f.name
        f.write(raw_html)

    print(f"Created temporary file: {temp_file}")

    try:
        # Now run style_injector
        print("Injecting styles and navigation...")
        style_injector.inject_styles_and_nav(
            temp_file,
            output_filename,
            layout_mode="sidebar",
            page_title=page_title,
            sidebar_title=sidebar_title
        )
    finally:
        # Cleanup
        if os.path.exists(temp_file):
            os.remove(temp_file)
            print("Cleaned up temp file.")

    if not os.path.exists(output_filename):
        print("Error: style injection failed, no output produced.")
        sys.exit(1)

    print(f"Done! Created {output_filename}")

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Convert DOCX to High-Fidelity HTML with Sidebar")
    parser.add_argument("input", help="Input .docx file path")
    parser.add_argument("output", nargs="?", help="Output .html file path (optional)")
    parser.add_argument("--title", default="教學手冊", help="Page title")
    parser.add_argument("--sidebar-title", default="教學手冊導航", help="Sidebar navigation title")

    args = parser.parse_args()
    
    output_path = args.output
    if not output_path:
        base_name = os.path.splitext(args.input)[0]
        output_path = f"{base_name}.html"
        
    convert_docx_to_html(args.input, output_path, page_title=args.title, sidebar_title=args.sidebar_title)
