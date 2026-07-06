"""
股票溝通師 — 完整 DOCX 報告產生器（V3）
=========================================

把 analyze_stock.py 的 JSON 結果 + yfinance 股價歷史，組成一份含 7 張圖的 Word 報告。
每張圖都附「怎麼看」教學說明 + 對應這次分析的具體解讀。

設計理念：
    - 終端機（Claude 對話端）= 即時對話，看文字摘要
    - DOCX 報告 = 完整紀錄、有圖、有教學，可保存可分享
    - 每張圖都是一個小教學 —— 工作坊用的時候老師可直接拿這個報告當教材

用法：
    python generate_report.py <股票代號> [中文別名清單]

範例：
    python generate_report.py 2454 "MTK,聯發科,聯發,MediaTek"

輸出：
    reports/<symbol>_<YYYYMMDD>_<HHMM>.docx
"""

from __future__ import annotations

import os
import sys
import warnings
from datetime import datetime
from pathlib import Path

import matplotlib
matplotlib.use("Agg")  # 不開視窗，純檔案輸出
import matplotlib.pyplot as plt
from matplotlib import font_manager
import numpy as np

warnings.filterwarnings("ignore")

# === 字型設定 ===
# 用 skill 內帶的 Noto Sans TC 字型，避免 Linux 系統沒中文字型。
SCRIPT_DIR = Path(__file__).resolve().parent
SKILL_DIR = SCRIPT_DIR.parent
FONT_PATH = SKILL_DIR / "assets" / "fonts" / "NotoSansTC.ttf"

if FONT_PATH.exists():
    # 把字型加進 matplotlib 字型管理員
    font_manager.fontManager.addfont(str(FONT_PATH))
    plt.rcParams["font.family"] = "Noto Sans TC"
    plt.rcParams["axes.unicode_minus"] = False  # 避免「-」變成方框
    CHINESE_FP = font_manager.FontProperties(fname=str(FONT_PATH))
else:
    CHINESE_FP = None
    print("⚠️ 找不到中文字型，圖表中文會變方框", file=sys.stderr)

# 把 analyze_stock 當模組 import（這樣不用透過 subprocess + JSON 接管道）
sys.path.insert(0, str(SCRIPT_DIR))
import analyze_stock  # noqa: E402

# 注意：這裡刻意「不」呼叫 analyze_stock._silence_stderr()。
# 那是給 analyze_stock.py 直接執行時保 stdout JSON 純度用的；
# 本腳本輸出的是 DOCX 檔不是 JSON，關掉 stderr 只會把
# SystemExit 訊息與 traceback 全部吞掉，失敗時變成無聲無息。


# ==============================================
# 圖表產生器（每張圖一個函式，畫完存到 tmp 路徑）
# ==============================================

CHART_STYLE = {
    "figure.facecolor": "white",
    "axes.facecolor": "#fafafa",
    "axes.edgecolor": "#333",
    "axes.grid": True,
    "grid.color": "#ddd",
    "grid.linestyle": "--",
    "grid.linewidth": 0.5,
    "axes.titleweight": "bold",
}


def _setup_fig(title: str, figsize=(9, 5)):
    """每張圖通用初始化。"""
    plt.style.use("default")
    for k, v in CHART_STYLE.items():
        plt.rcParams[k] = v
    fig, ax = plt.subplots(figsize=figsize)
    ax.set_title(title, fontsize=14, pad=12, fontproperties=CHINESE_FP)
    return fig, ax


def _save_and_close(fig, out_path: Path):
    fig.tight_layout()
    fig.savefig(str(out_path), dpi=120, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def chart_kline_with_ma(hist, symbol: str, name: str, out_path: Path):
    """圖 1：K 線 + 5/20/60 日均線。"""
    fig, ax = _setup_fig(f"{symbol} {name} — 近 6 個月日 K 與三均線")

    close = hist["Close"]
    high = hist["High"]
    low = hist["Low"]
    open_ = hist["Open"]
    dates = hist.index

    # K 線：上漲日紅、下跌日綠（台股慣例）
    for i, (o, h, l, c, d) in enumerate(zip(open_, high, low, close, dates)):
        color = "#d63031" if c >= o else "#00b894"
        ax.vlines(i, l, h, color=color, linewidth=0.8)
        ax.add_patch(plt.Rectangle((i - 0.3, min(o, c)), 0.6, abs(c - o) or 0.01,
                                    color=color, alpha=0.85))

    # 三均線
    ma5 = close.rolling(5).mean()
    ma20 = close.rolling(20).mean()
    ma60 = close.rolling(60).mean()
    x = range(len(close))
    ax.plot(x, ma5, label="5 日均線", color="#fdcb6e", linewidth=1.5)
    ax.plot(x, ma20, label="20 日均線", color="#0984e3", linewidth=1.5)
    ax.plot(x, ma60, label="60 日均線", color="#6c5ce7", linewidth=1.8)

    # x 軸日期格式
    n = len(dates)
    step = max(n // 8, 1)
    ax.set_xticks(range(0, n, step))
    ax.set_xticklabels([dates[i].strftime("%m/%d") for i in range(0, n, step)],
                        rotation=0, fontsize=9)
    ax.set_ylabel("股價", fontproperties=CHINESE_FP)
    ax.legend(loc="upper left", prop=CHINESE_FP, fontsize=9)
    _save_and_close(fig, out_path)


def chart_volume(hist, symbol: str, out_path: Path):
    """圖 2：成交量柱狀圖 + 20 日均量水平線。"""
    fig, ax = _setup_fig(f"{symbol} — 成交量 vs 20 日均量")
    volume = hist["Volume"]
    close = hist["Close"]
    open_ = hist["Open"]

    # 上漲日紅、下跌日綠
    colors = ["#d63031" if c >= o else "#00b894" for c, o in zip(close, open_)]
    x = range(len(volume))
    ax.bar(x, volume, color=colors, width=0.7, alpha=0.85)

    # 20 日均量水平線
    avg_20 = volume.rolling(20).mean()
    ax.plot(x, avg_20, color="#2d3436", linewidth=1.5, label="20 日均量", linestyle="--")

    n = len(volume)
    step = max(n // 8, 1)
    ax.set_xticks(range(0, n, step))
    ax.set_xticklabels([hist.index[i].strftime("%m/%d") for i in range(0, n, step)],
                        rotation=0, fontsize=9)
    ax.set_ylabel("成交量（股）", fontproperties=CHINESE_FP)
    ax.legend(loc="upper left", prop=CHINESE_FP, fontsize=9)

    # y 軸用千分位
    ax.yaxis.set_major_formatter(plt.FuncFormatter(lambda x, _: f"{int(x):,}"))
    _save_and_close(fig, out_path)


def chart_rsi(hist, symbol: str, out_path: Path):
    """圖 3：RSI(14) 折線圖 + 70/30 警戒線。"""
    fig, ax = _setup_fig(f"{symbol} — RSI(14) 動能指標")

    close = hist["Close"]
    delta = close.diff()
    gain = delta.where(delta > 0, 0).rolling(14).mean()
    loss = (-delta.where(delta < 0, 0)).rolling(14).mean()
    rs = gain / loss.replace(0, 1e-10)
    rsi = 100 - (100 / (1 + rs))

    x = range(len(rsi))
    ax.plot(x, rsi, color="#e17055", linewidth=2)
    ax.axhline(70, color="#d63031", linestyle="--", linewidth=1, label="超買線（70）")
    ax.axhline(30, color="#00b894", linestyle="--", linewidth=1, label="超賣線（30）")
    ax.fill_between(x, 70, 100, color="#d63031", alpha=0.08)
    ax.fill_between(x, 0, 30, color="#00b894", alpha=0.08)
    ax.set_ylim(0, 100)

    n = len(rsi)
    step = max(n // 8, 1)
    ax.set_xticks(range(0, n, step))
    ax.set_xticklabels([hist.index[i].strftime("%m/%d") for i in range(0, n, step)],
                        rotation=0, fontsize=9)
    ax.set_ylabel("RSI", fontproperties=CHINESE_FP)
    ax.legend(loc="upper left", prop=CHINESE_FP, fontsize=9)
    _save_and_close(fig, out_path)


def chart_macd(hist, symbol: str, out_path: Path):
    """圖 4：MACD 三件套（DIF/DEA/柱）。"""
    fig, ax = _setup_fig(f"{symbol} — MACD 動能交叉")

    close = hist["Close"]
    ema12 = close.ewm(span=12, adjust=False).mean()
    ema26 = close.ewm(span=26, adjust=False).mean()
    dif = ema12 - ema26
    dea = dif.ewm(span=9, adjust=False).mean()
    hist_macd = dif - dea

    x = range(len(dif))
    # 柱：正紅、負綠
    colors = ["#d63031" if v >= 0 else "#00b894" for v in hist_macd]
    ax.bar(x, hist_macd, color=colors, width=0.7, alpha=0.7, label="MACD 柱")
    ax.plot(x, dif, color="#0984e3", linewidth=1.6, label="DIF")
    ax.plot(x, dea, color="#fdcb6e", linewidth=1.6, label="DEA")
    ax.axhline(0, color="#333", linewidth=0.6)

    n = len(dif)
    step = max(n // 8, 1)
    ax.set_xticks(range(0, n, step))
    ax.set_xticklabels([hist.index[i].strftime("%m/%d") for i in range(0, n, step)],
                        rotation=0, fontsize=9)
    ax.legend(loc="upper left", prop=CHINESE_FP, fontsize=9)
    _save_and_close(fig, out_path)


def chart_bollinger(hist, symbol: str, out_path: Path):
    """圖 5：布林通道 + 股價。"""
    fig, ax = _setup_fig(f"{symbol} — 布林通道（20MA ± 2σ）")
    close = hist["Close"]
    mid = close.rolling(20).mean()
    std = close.rolling(20).std()
    upper = mid + 2 * std
    lower = mid - 2 * std

    x = range(len(close))
    ax.plot(x, close, color="#2d3436", linewidth=1.5, label="股價")
    ax.plot(x, mid, color="#0984e3", linewidth=1.2, label="中軸（20MA）", linestyle="--")
    ax.plot(x, upper, color="#d63031", linewidth=1.2, label="上軌（+2σ）")
    ax.plot(x, lower, color="#00b894", linewidth=1.2, label="下軌（-2σ）")
    ax.fill_between(x, lower, upper, color="#0984e3", alpha=0.05)

    n = len(close)
    step = max(n // 8, 1)
    ax.set_xticks(range(0, n, step))
    ax.set_xticklabels([hist.index[i].strftime("%m/%d") for i in range(0, n, step)],
                        rotation=0, fontsize=9)
    ax.set_ylabel("股價", fontproperties=CHINESE_FP)
    ax.legend(loc="upper left", prop=CHINESE_FP, fontsize=9)
    _save_and_close(fig, out_path)


def chart_institutional(institutional: dict, symbol: str, out_path: Path) -> bool:
    """圖 6：三大法人單日買賣超分組柱狀圖。回傳 True 表示有產出。"""
    if not institutional:
        return False
    fig, ax = _setup_fig(f"{symbol} — 三大法人買賣超（{institutional['trade_date']}）")

    categories = ["外資", "投信", "自營商", "三大法人合計"]
    values = [
        institutional.get("foreign_net", 0),
        institutional.get("trust_net", 0),
        institutional.get("dealer_net", 0),
        institutional.get("total_net", 0),
    ]
    # 化為「萬股」較直觀（1 萬股 = 10000）
    values_wan = [v / 10000 for v in values]
    colors = ["#d63031" if v >= 0 else "#00b894" for v in values]

    bars = ax.bar(categories, values_wan, color=colors, alpha=0.85, width=0.6)
    ax.axhline(0, color="#333", linewidth=0.8)

    # 在柱上標數字
    max_abs = max(abs(v) for v in values_wan) or 1
    for bar, v in zip(bars, values_wan):
        h = bar.get_height()
        offset = max_abs * 0.02
        ax.text(bar.get_x() + bar.get_width() / 2,
                h + (offset if h >= 0 else -offset),
                f"{v:+,.0f}",
                ha="center",
                va="bottom" if h >= 0 else "top",
                fontproperties=CHINESE_FP, fontsize=10)

    ax.set_ylabel("買賣超（萬股）", fontproperties=CHINESE_FP)
    plt.setp(ax.get_xticklabels(), fontproperties=CHINESE_FP, fontsize=11)
    _save_and_close(fig, out_path)
    return True


def chart_heat_radar(heat_summary: dict, ptt: dict | None, yahoo: dict | None,
                     news_density: dict | None, volume: dict, institutional: dict | None,
                     symbol: str, out_path: Path) -> bool:
    """圖 7：多維熱度雷達圖（4 軸）。"""
    if not heat_summary:
        return False

    # 把每個維度的「熱度等級」量化到 0-5 分
    def score_ptt(p):
        if not p: return 0
        rate = p.get("mention_rate_pct", 0)
        if rate >= 10: return 5
        if rate >= 5: return 4
        if rate >= 3: return 3
        if rate >= 1: return 2
        if rate > 0: return 1
        return 0

    def score_news(nd):
        if not nd: return 0
        if nd.get("ratio") is not None:
            r = nd["ratio"]
            if r >= 3: return 5
            if r >= 1.5: return 4
            if r >= 0.7: return 3
            if r > 0: return 2
            return 1
        # 沒有歷史對比時用近 3 日則數
        c = nd.get("recent_3d_count", 0)
        if c >= 8: return 5
        if c >= 5: return 4
        if c >= 2: return 3
        if c >= 1: return 2
        return 0

    def score_volume(v):
        ratio = v.get("ratio_to_20d") or 0
        if ratio >= 2: return 5
        if ratio >= 1.5: return 4
        if ratio >= 1: return 3
        if ratio >= 0.7: return 2
        if ratio > 0.3: return 1
        return 0

    def score_institutional(i):
        if not i: return 0
        net = abs(i.get("total_net", 0))
        # 以股數分級（1000 萬以上極大、500 萬大、100 萬中等）
        if net >= 10_000_000: return 5
        if net >= 5_000_000: return 4
        if net >= 1_000_000: return 3
        if net >= 100_000: return 2
        if net > 0: return 1
        return 0

    categories = ["散戶熱度\n(PTT)", "媒體熱度\n(新聞密度)",
                  "市場熱度\n(量比)", "法人方向\n(買賣超強度)"]
    scores = [
        score_ptt(ptt),
        score_news(news_density),
        score_volume(volume),
        score_institutional(institutional),
    ]

    # 雷達圖
    angles = np.linspace(0, 2 * np.pi, len(categories), endpoint=False).tolist()
    scores_closed = scores + [scores[0]]
    angles_closed = angles + [angles[0]]

    fig, ax = plt.subplots(figsize=(7, 7), subplot_kw=dict(polar=True))
    ax.fill(angles_closed, scores_closed, color="#0984e3", alpha=0.25)
    ax.plot(angles_closed, scores_closed, color="#0984e3", linewidth=2)
    ax.set_ylim(0, 5)
    ax.set_yticks([1, 2, 3, 4, 5])
    ax.set_yticklabels(["1\n微弱", "2\n偏弱", "3\n中等", "4\n偏強", "5\n極強"],
                       fontsize=8, fontproperties=CHINESE_FP)
    ax.set_xticks(angles)
    ax.set_xticklabels(categories, fontproperties=CHINESE_FP, fontsize=11)
    ax.set_title(f"{symbol} — 四維熱度雷達",
                 fontsize=14, pad=20, fontproperties=CHINESE_FP)

    # 在每個頂點標分數
    for ang, s, cat in zip(angles, scores, categories):
        ax.text(ang, s + 0.3, f"{s}",
                ha="center", va="center",
                fontproperties=CHINESE_FP,
                fontweight="bold", fontsize=12, color="#0984e3")

    plt.tight_layout()
    fig.savefig(str(out_path), dpi=120, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return True


# ==============================================
# DOCX 報告組裝
# ==============================================

# 每張圖的「怎麼看」教學說明（固定文字 + 動態解讀）
CHART_GUIDES = {
    "kline": {
        "title": "📊 圖 1：K 線 + 三均線疊圖",
        "how_to_read": (
            "**怎麼看這張圖：**\n\n"
            "- **K 線顏色**：紅色代表當日上漲（收盤 ≥ 開盤）；綠色代表當日下跌。\n"
            "- **三條均線**：黃線（5 日）= 短期、藍線（20 日）= 中期、紫線（60 日）= 長期。\n"
            "- **多頭排列** = 5 日線在上、20 日線中、60 日線下，趨勢健康。\n"
            "- **空頭排列** = 三線順序倒過來，趨勢偏弱。\n"
            "- **三線糾結** = 三條線交織，方向不明朗（盤整中）。\n"
            "- **股價跌破 60 日線**通常被視為中期支撐失守，要特別留意。"
        ),
    },
    "volume": {
        "title": "📊 圖 2：成交量 vs 20 日均量",
        "how_to_read": (
            "**怎麼看這張圖：**\n\n"
            "- **柱子顏色**：紅色 = 當日上漲、綠色 = 當日下跌。\n"
            "- **黑色虛線**：20 日均量（最近一個月平均成交量）。\n"
            "- **量柱超過虛線 1.5 倍** = 「爆量」，常出現在重大消息或主力進出時。\n"
            "- **連續多日量柱在虛線下方** = 量縮觀望，市場沒共識。\n"
            "- **量價組合**：量增價漲（健康）、量縮價漲（小心動能不足）、量增價跌（賣壓重）、量縮價跌（恐慌可能尾聲）。"
        ),
    },
    "rsi": {
        "title": "📊 圖 3：RSI(14) 動能指標",
        "how_to_read": (
            "**怎麼看這張圖：**\n\n"
            "- **RSI 是什麼**：衡量「最近上漲動能 vs 下跌動能」的比例，0-100 之間。\n"
            "- **紅色虛線（70）**：超買警戒線，超過代表「短期過熱」。\n"
            "- **綠色虛線（30）**：超賣警戒線，低於代表「短期超賣」。\n"
            "- **重點不是穿越瞬間，而是穿越後 5 天往哪走**：RSI 從 80 往下彎才是真的拉回，硬撐在 70 以上不一定立即反轉。\n"
            "- **背離訊號**：股價創新高但 RSI 沒創新高 = 動能不足，要警戒。"
        ),
    },
    "macd": {
        "title": "📊 圖 4：MACD 動能交叉",
        "how_to_read": (
            "**怎麼看這張圖：**\n\n"
            "- **DIF（藍線）**：快慢線差，反映短期動能。\n"
            "- **DEA（橘線）**：DIF 的平滑線，作為訊號基準。\n"
            "- **柱（紅綠）**：DIF − DEA 的差距視覺化。柱越長 = 動能越強。\n"
            "- **黃金交叉**：DIF 由下穿上 DEA → 多方訊號開始。\n"
            "- **死亡交叉**：DIF 由上穿下 DEA → 空方訊號開始。\n"
            "- **零軸**：DIF/DEA 在零軸上方代表多方優勢、下方代表空方優勢。"
        ),
    },
    "bollinger": {
        "title": "📊 圖 5：布林通道",
        "how_to_read": (
            "**怎麼看這張圖：**\n\n"
            "- **中軸（藍虛線）**：20 日均線，是布林通道的中心。\n"
            "- **上軌（紅實線）**：中軸 + 2 倍標準差。\n"
            "- **下軌（綠實線）**：中軸 − 2 倍標準差。\n"
            "- **股價貼上軌**：短期過熱，常見於強勢股的飆漲段。\n"
            "- **股價貼下軌**：短期超賣，可能是反彈契機。\n"
            "- **通道收斂**（上下軌靠近）= 波動率變低，常預告下一波大行情。\n"
            "- **通道擴張** = 波動加劇，趨勢明顯。"
        ),
    },
    "institutional": {
        "title": "📊 圖 6：三大法人買賣超（台股限定）",
        "how_to_read": (
            "**怎麼看這張圖：**\n\n"
            "- **外資**：通常是大頭，買賣量最大。\n"
            "- **投信**：本土基金，中等規模，常被視為「在地觀點」。\n"
            "- **自營商**：券商自己的部位，偏短線。\n"
            "- **三大法人合計**：上面三者加總，是最常被引用的「法人總方向」。\n"
            "- **三條柱都向上** = 三大法人都在買 → 強烈買盤共識。\n"
            "- **三條柱都向下** = 三大法人都在賣 → 強烈賣壓共識。\n"
            "- **方向相反**（有人買有人賣）= 法人意見分歧，散戶要更小心解讀。"
        ),
    },
    "heat_radar": {
        "title": "📊 圖 7：四維熱度雷達",
        "how_to_read": (
            "**怎麼看這張圖：**\n\n"
            "- **四個方向**：散戶熱度（PTT）、媒體熱度（新聞密度）、市場熱度（量比）、法人方向（買賣超強度）。\n"
            "- **每個軸 0-5 分**：分數越外緣代表該維度訊號越強。\n"
            "- **四個方向都很強** → 高度共識（多空雙方都熱絡）。\n"
            "- **僅一個方向特別突出** → 有人跟其他人意見不同，是值得追究的線索。\n"
            "- **散戶熱（PTT 滿格）+ 法人弱（買賣超少）** → 「散戶在追、法人觀望」的危險訊號。\n"
            "- **法人強（買賣超大）+ 散戶冷** → 「法人悄悄佈局或出貨」的線索，看法人方向是買還是賣判斷。"
        ),
    },
}


def add_section(doc, heading: str, level: int = 1):
    h = doc.add_heading(heading, level=level)
    return h


def add_image_with_caption(doc, img_path: Path, guide_key: str, dynamic_reading: str | None = None):
    """放圖 + 怎麼看說明 + 這次解讀。"""
    from docx.shared import Inches
    guide = CHART_GUIDES[guide_key]
    doc.add_heading(guide["title"], level=2)
    doc.add_picture(str(img_path), width=Inches(6.0))

    # 怎麼看（教學）
    p_how = doc.add_paragraph()
    p_how.add_run("📖 怎麼看：").bold = True
    for line in guide["how_to_read"].split("\n"):
        if line.strip():
            stripped = line.strip()
            # 整行被 ** 包住的視為小標、套 Intense Quote
            is_heading = stripped.startswith("**") and stripped.endswith("**")
            p = doc.add_paragraph(style="Intense Quote" if is_heading else None)
            _add_md_runs(p, stripped)

    # 這次的具體解讀
    if dynamic_reading:
        p_dyn = doc.add_paragraph()
        p_dyn.add_run("🔍 這次的解讀：").bold = True
        for line in dynamic_reading.split("\n"):
            if line.strip():
                p = doc.add_paragraph()
                _add_md_runs(p, line.strip())


def build_dynamic_kline_reading(d: dict) -> str:
    trend = d.get("trend", {})
    notes = []
    notes.append(f"目前狀態：{trend.get('label', '—')}（5MA={trend.get('ma5')}、20MA={trend.get('ma20')}、60MA={trend.get('ma60')}）。")
    above = sum(1 for k in ["above_ma5", "above_ma20", "above_ma60"] if trend.get(k))
    notes.append(f"股價站上 {above} 條均線（共 3 條）。")
    pos = d.get("price", {}).get("position_in_60d_range_pct")
    if pos is not None:
        notes.append(f"位於近 60 日區間 {pos}% 位置（高點={d['price']['high_60d']}、低點={d['price']['low_60d']}）。")
    return "\n".join(notes)


def build_dynamic_volume_reading(d: dict) -> str:
    v = d.get("volume", {})
    notes = [f"量比 20MA：{v.get('ratio_to_20d')} 倍 → {v.get('label')}"]
    if v.get("intraday_partial_warning"):
        notes.append("⚠️ 注意：當日量可能尚未累積完畢（盤中），數值僅供參考。")
    return "\n".join(notes)


def build_dynamic_rsi_reading(d: dict) -> str:
    rsi = d.get("momentum", {}).get("rsi14")
    if rsi is None:
        return "RSI 資料不足。"
    if rsi >= 80:
        zone = "極度超買區（>80），短期拉回風險高"
    elif rsi >= 70:
        zone = "超買區（70-80），偏熱"
    elif rsi >= 30:
        zone = "中性區（30-70）"
    elif rsi >= 20:
        zone = "超賣區（20-30），偏冷"
    else:
        zone = "極度超賣區（<20），可能有反彈契機"
    return f"目前 RSI = {rsi}，落在 {zone}。"


def build_dynamic_macd_reading(d: dict) -> str:
    m = d.get("momentum", {})
    return f"DIF = {m.get('macd_dif')}、DEA = {m.get('macd_dea')} → {m.get('macd_label', '—')}"


def build_dynamic_bollinger_reading(d: dict) -> str:
    v = d.get("volatility", {})
    notes = [f"目前布林位置：{v.get('bollinger_label', '—')}"]
    if v.get("bb_upper") and v.get("bb_lower"):
        notes.append(f"上軌約 {v['bb_upper']}、下軌約 {v['bb_lower']}。")
    if v.get("annualized_pct") is not None:
        notes.append(f"年化波動率：{v['annualized_pct']}%。")
    return "\n".join(notes)


def build_dynamic_inst_reading(institutional: dict) -> str:
    return (
        f"日期：{institutional['trade_date']}\n"
        f"外資：{institutional['foreign_label']}\n"
        f"投信：{institutional['trust_label']}\n"
        f"自營商：{institutional['dealer_label']}\n"
        f"合計：{institutional['total_label']}（{institutional['stance']}）"
    )


def build_dynamic_radar_reading(heat_summary: dict) -> str:
    notes = []
    if heat_summary.get("retail_heat"):
        notes.append("散戶熱度訊號：" + "；".join(heat_summary["retail_heat"]))
    if heat_summary.get("media_heat"):
        notes.append(f"媒體熱度：{heat_summary['media_heat']}")
    if heat_summary.get("market_heat"):
        notes.append(f"市場熱度：{heat_summary['market_heat']}")
    if heat_summary.get("institutional_stance"):
        notes.append(f"法人方向：{heat_summary['institutional_stance']}")
    cc = heat_summary.get("cross_check", [])
    if cc:
        notes.append("\n交叉判讀：")
        notes.extend(["• " + c for c in cc])
    return "\n".join(notes)


def build_auto_persona(data: dict) -> str:
    """沒提供 --persona 時，用模板自動生成擬人化獨白（堪用版本）。

    句型結構：開場（我幫你問了 X 哦）→ 描述狀態 → 揭露隱憂或得意 → 給提醒
    """
    name = data.get("name", "") or data.get("symbol", "")
    short_name = data.get("symbol", "")
    r60 = data.get("price", {}).get("return_60d_pct") or 0
    rsi = data.get("momentum", {}).get("rsi14") or 50
    trend = data.get("trend", {}).get("label", "")
    inst = data.get("institutional")
    inst_stance = inst.get("stance", "") if inst else ""
    warnings_count = len(data.get("warnings", []))
    pos_pct = data.get("price", {}).get("position_in_60d_range_pct") or 50

    # 句 1：當前狀態
    if r60 > 30 and rsi > 70:
        s1 = f"他說他最近真的飛得很高，60 天漲了 {r60:.0f}%，每天都覺得有點頭暈"
    elif r60 > 10:
        s1 = f"他說最近狀態不錯，60 天累積漲了 {r60:.0f}%，腳步算穩"
    elif r60 < -15:
        s1 = f"他抱怨最近真的很不順，60 天跌了 {abs(r60):.0f}%，連喘息的空間都沒有"
    elif r60 < -5:
        s1 = f"他說最近有點低潮，60 天跌了 {abs(r60):.0f}%，提不起勁"
    else:
        s1 = f"他說最近沒什麼大事，60 天大致持平（{r60:+.1f}%），就是在原地踱步"

    # 句 2：隱憂或得意（基於法人和量價）
    if "賣超" in inst_stance:
        s2 = f"他偷偷跟我說，{inst.get('total_label', '法人在減碼')}，他自己也擔心這個訊號"
    elif "買超" in inst_stance:
        s2 = f"他有點得意地說，{inst.get('total_label', '法人在加碼')}，他覺得有人挺他"
    else:
        s2 = "他說目前沒有特別劇烈的法人動作，市場還在觀望"

    # 句 3：提醒
    if warnings_count >= 2:
        s3 = f"他要我提醒你：『現在訊號上有 {warnings_count} 條警訊亮燈，你要小心一點』"
    elif pos_pct > 90:
        s3 = "他提醒：『我現在站在 60 日的高點區，追進來的人請有準備被甩下去』"
    elif pos_pct < 20:
        s3 = "他提醒：『我在低點區徘徊，可能還沒見底，但也可能在醞釀反彈』"
    else:
        s3 = f"他要我跟你說：『目前還算淡定，但記得我是 {trend}，狀態會變』"

    return (
        f"我幫你問了 {short_name} 哦 ——\n\n"
        f"{s1}。\n\n"
        f"{s2}。\n\n"
        f"{s3}"
    )


def build_auto_summary(data: dict) -> str:
    """沒提供 --summary 時，用模板自動生成一句話總結。"""
    rsi = data.get("momentum", {}).get("rsi14") or 50
    r60 = data.get("price", {}).get("return_60d_pct") or 0
    trend = data.get("trend", {}).get("label", "")
    inst = data.get("institutional")
    warnings_count = len(data.get("warnings", []))

    # 風險評估
    risk_signals = []
    if rsi > 80:
        risk_signals.append(f"RSI {rsi:.0f} 極度超買")
    elif rsi > 70:
        risk_signals.append(f"RSI {rsi:.0f} 偏熱")
    if r60 > 50:
        risk_signals.append(f"60 日暴漲 {r60:.0f}%")
    if inst and inst.get("total_net", 0) < -1_000_000:
        risk_signals.append(f"法人 {inst.get('total_label', '賣超')}")
    if warnings_count > 0:
        risk_signals.append(f"{warnings_count} 條技術警訊")

    if len(risk_signals) >= 2:
        return (
            f"目前訊號上偏熱：{'、'.join(risk_signals)}。"
            f"短線追高風險高，建議等回檔或法人轉買再考慮分批進場。"
            f"⚠️ 本分析僅為技術面解讀與娛樂用途，非投資建議。"
        )

    if rsi < 30 and r60 < -10:
        return (
            f"目前訊號上偏冷：RSI {rsi:.0f} 超賣、60 日跌 {abs(r60):.0f}%。"
            f"短線可能有反彈機會，但要看法人態度與技術止穩訊號再分批佈局。"
            f"⚠️ 本分析僅為技術面解讀與娛樂用途，非投資建議。"
        )

    return (
        f"目前處於相對中性的 {trend} 狀態，60 日報酬 {r60:+.1f}%、RSI {rsi:.0f}。"
        f"沒有特別強烈的方向訊號，建議持續觀察並結合自己的進出邏輯。"
        f"⚠️ 本分析僅為技術面解讀與娛樂用途，非投資建議。"
    )


def _md_segments(text: str) -> list[tuple[str, bool]]:
    """把含 **bold** 標記的 Markdown 字串切成 (text, is_bold) 片段。

    例如 '今天 **跌破 60 日線**，要小心' →
        [('今天 ', False), ('跌破 60 日線', True), ('，要小心', False)]

    沒 ** 標記就回 [(text, False)]。落單的 ** 視為純文字（不嘗試聰明配對）。
    """
    import re
    segments: list[tuple[str, bool]] = []
    pattern = re.compile(r"\*\*(.+?)\*\*")
    last = 0
    for m in pattern.finditer(text):
        if m.start() > last:
            segments.append((text[last:m.start()], False))
        segments.append((m.group(1), True))
        last = m.end()
    if last < len(text):
        segments.append((text[last:], False))
    return segments if segments else [(text, False)]


def _add_md_runs(paragraph, text: str, base_bold: bool = False,
                  base_italic: bool = False, base_color=None,
                  base_size=None) -> None:
    """把含 **bold** 的字串拆成多個 run 加進 paragraph，保留 base 樣式。

    base_bold / base_italic / base_color / base_size 是套在所有 run 上的基底樣式；
    **xxx** 片段會在基底之上**疊加** bold（即使 base_bold=False 也會變粗）。
    """
    for seg_text, seg_bold in _md_segments(text):
        run = paragraph.add_run(seg_text)
        run.bold = base_bold or seg_bold
        if base_italic:
            run.italic = True
        if base_color is not None:
            run.font.color.rgb = base_color
        if base_size is not None:
            run.font.size = base_size


def add_blockquote_paragraph(doc, text: str):
    """加入引言塊樣式的段落（用於擬人化獨白）。支援 **bold** Markdown。"""
    from docx.shared import RGBColor, Pt
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(20)
    _add_md_runs(p, text, base_italic=True, base_color=RGBColor(0x55, 0x55, 0x55))


def build_report(symbol: str, name_zh: str | None,
                  persona: str | None = None, summary: str | None = None,
                  output_dir: Path | str | None = None) -> Path:
    """主流程：抓資料 → 畫 7 張圖 → 組 DOCX → 回傳檔案路徑。

    Args:
        symbol: 股票代號
        name_zh: 中文別名清單（逗號分隔）
        persona: Claude 在對話中產生的擬人化獨白；None 時用模板自動生成
        summary: Claude 在對話中產生的一句話總結；None 時用模板自動生成
        output_dir: 報告輸出目錄；None 時用「呼叫端 cwd / reports」。
                    刻意不寫進 skill 資料夾，避免使用者報告污染共享 skill 目錄。
    """
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    print(f"➜ 抓 {symbol} 完整資料中...")
    data = analyze_stock.analyze(symbol, name_zh)
    if "error" in data:
        raise SystemExit(f"❌ 分析失敗：{data['error']}")

    # 重新抓股價歷史用來畫圖（analyze() 內部沒回傳 DataFrame）
    print("➜ 抓股價歷史中...")
    hist, info, err = analyze_stock.fetch_price_history(data["symbol"])
    if err or hist is None or hist.empty:
        raise SystemExit(f"❌ 抓股價失敗：{err}")

    # 畫圖到 tmp 路徑
    print("➜ 畫圖中（7 張）...")
    timestamp = datetime.now().strftime("%Y%m%d_%H%M")
    if output_dir is None:
        reports_dir = Path.cwd() / "reports"
    else:
        reports_dir = Path(output_dir)
    reports_dir.mkdir(parents=True, exist_ok=True)
    tmp_dir = reports_dir / f".tmp_{timestamp}"
    tmp_dir.mkdir(exist_ok=True)

    # V3.4：畫圖～組 DOCX 整段包 try/finally——就算中途丟例外（yfinance 斷線、
    # matplotlib 字型問題等），也不在使用者目錄留下 .tmp_* 暫存圖目錄。
    try:
        bare = data["symbol"].replace(".", "_")
        chart_paths = {}

        chart_kline_with_ma(hist, data["symbol"], data.get("name", ""),
                             tmp_dir / f"{bare}_1_kline.png")
        chart_paths["kline"] = tmp_dir / f"{bare}_1_kline.png"

        chart_volume(hist, data["symbol"], tmp_dir / f"{bare}_2_volume.png")
        chart_paths["volume"] = tmp_dir / f"{bare}_2_volume.png"

        chart_rsi(hist, data["symbol"], tmp_dir / f"{bare}_3_rsi.png")
        chart_paths["rsi"] = tmp_dir / f"{bare}_3_rsi.png"

        chart_macd(hist, data["symbol"], tmp_dir / f"{bare}_4_macd.png")
        chart_paths["macd"] = tmp_dir / f"{bare}_4_macd.png"

        chart_bollinger(hist, data["symbol"], tmp_dir / f"{bare}_5_bollinger.png")
        chart_paths["bollinger"] = tmp_dir / f"{bare}_5_bollinger.png"

        has_inst = chart_institutional(data.get("institutional"), data["symbol"],
                                        tmp_dir / f"{bare}_6_institutional.png")
        if has_inst:
            chart_paths["institutional"] = tmp_dir / f"{bare}_6_institutional.png"

        has_radar = chart_heat_radar(
            data.get("heat_summary"), data.get("ptt_heat"), data.get("yahoo_community"),
            data.get("news_density"), data.get("volume", {}), data.get("institutional"),
            data["symbol"], tmp_dir / f"{bare}_7_radar.png"
        )
        if has_radar:
            chart_paths["heat_radar"] = tmp_dir / f"{bare}_7_radar.png"

        # === 組 DOCX ===
        print("➜ 組 DOCX 中...")
        doc = Document()

        # 預設字型
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

        # 封面
        title = doc.add_heading(f"{data['symbol']} {data.get('name', '')}", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub.add_run("股票溝通師完整分析報告\n").bold = True
        sub.add_run(f"分析時間：{data.get('analyzed_at', '')}\n")
        sub.add_run(f"市場：{data.get('market', '')}　產業：{data.get('sector', '')}")

        # 免責聲明
        discl = doc.add_paragraph()
        discl.alignment = WD_ALIGN_PARAGRAPH.CENTER
        discl_run = discl.add_run(
            "⚠️ 本報告僅為技術面解讀與娛樂用途，非投資建議。投資決策請自行評估並承擔風險。"
        )
        discl_run.italic = True
        discl_run.font.color.rgb = RGBColor(0xCC, 0x44, 0x44)

        # === 🎤 他想說什麼（擬人化獨白，緊接封面後）===
        # 設計哲學：「股票溝通師」先讓股票自己說話，再給技術分析
        add_section(doc, "🎤 他想說什麼（先聽他講話）")

        persona_text = persona if persona else build_auto_persona(data)
        # 把擬人化獨白以引言塊樣式呈現，每段一個 paragraph
        for line in persona_text.split("\n"):
            line = line.strip()
            if line:
                add_blockquote_paragraph(doc, line)

        if not persona:
            # 自動生成的話加個小註記，讓使用者知道（這在工作坊上有教學意義）
            note = doc.add_paragraph()
            note.add_run("（以上獨白由模板自動生成；對話中的客製化版本可能更貼合你的處境）").italic = True
            from docx.shared import Pt as _Pt
            note.runs[0].font.size = _Pt(9)

        # === 📌 一頁摘要 ===
        add_section(doc, "📌 一頁摘要")

        p = doc.add_paragraph()
        p.add_run(f"最新價：").bold = True
        p.add_run(f"{data['price']['latest_close']}")
        p.add_run("　│　近 60 日報酬：").bold = True
        p.add_run(f"{data['price'].get('return_60d_pct')}%")
        p.add_run("　│　趨勢：").bold = True
        p.add_run(f"{data['trend']['label']}")
        p.add_run("　│　RSI：").bold = True
        p.add_run(f"{data['momentum']['rsi14']}")

        # 警訊
        warnings_list = data.get("warnings", [])
        if warnings_list:
            doc.add_heading("⚠️ 警訊清單", level=2)
            for w in warnings_list:
                doc.add_paragraph(w, style="List Bullet")
        else:
            doc.add_paragraph("✅ 目前沒有觸發任何技術警訊。")

        # 客觀觀察
        notes = data.get("notes", [])
        if notes:
            doc.add_heading("🔬 客觀觀察", level=2)
            for n in notes:
                p = doc.add_paragraph(style="List Bullet")
                _add_md_runs(p, n)

        # 多維熱度交叉判讀
        hs = data.get("heat_summary", {})
        cc = hs.get("cross_check", [])
        if cc:
            doc.add_heading("📊 多維熱度交叉判讀", level=2)
            for c in cc:
                p = doc.add_paragraph(style="List Bullet")
                _add_md_runs(p, c)

        # === 圖表區塊 ===
        doc.add_page_break()
        add_section(doc, "📊 技術圖表逐一解讀")

        intro = doc.add_paragraph()
        intro.add_run(
            "下面 7 張圖是這次分析的視覺化呈現。每張圖都附「怎麼看」教學說明（適用所有股票）"
            "和「這次的解讀」（針對此次標的的具體狀況）。"
        )

        add_image_with_caption(doc, chart_paths["kline"], "kline",
                               build_dynamic_kline_reading(data))
        add_image_with_caption(doc, chart_paths["volume"], "volume",
                               build_dynamic_volume_reading(data))
        add_image_with_caption(doc, chart_paths["rsi"], "rsi",
                               build_dynamic_rsi_reading(data))
        add_image_with_caption(doc, chart_paths["macd"], "macd",
                               build_dynamic_macd_reading(data))
        add_image_with_caption(doc, chart_paths["bollinger"], "bollinger",
                               build_dynamic_bollinger_reading(data))

        if "institutional" in chart_paths:
            add_image_with_caption(doc, chart_paths["institutional"], "institutional",
                                   build_dynamic_inst_reading(data["institutional"]))
        else:
            doc.add_paragraph("💡 此標的非台股，無三大法人資料。")

        if "heat_radar" in chart_paths:
            add_image_with_caption(doc, chart_paths["heat_radar"], "heat_radar",
                                   build_dynamic_radar_reading(data.get("heat_summary", {})))

        # === 📌 一句話總結（必要結尾）===
        doc.add_page_break()
        add_section(doc, "📌 一句話總結")

        summary_text = summary if summary else build_auto_summary(data)
        summary_p = doc.add_paragraph()
        # 整段預設 bold + 12pt；內含 **xxx** 也不會疊加（已是 bold）
        _add_md_runs(summary_p, summary_text, base_bold=True, base_size=Pt(12))

        if not summary:
            # 自動生成註記
            note = doc.add_paragraph()
            note.add_run("（以上總結由模板自動生成；對話中的客製化版本會根據你的持倉/時間/金額更精準）").italic = True
            note.runs[0].font.size = Pt(9)

        # 結尾免責
        final_discl = doc.add_paragraph()
        final_discl.alignment = WD_ALIGN_PARAGRAPH.CENTER
        final_discl_run = final_discl.add_run(
            "⚠️ 本分析僅為技術面解讀與娛樂用途，非投資建議。投資決策請自行評估並承擔風險。"
        )
        final_discl_run.italic = True
        final_discl_run.font.color.rgb = RGBColor(0xCC, 0x44, 0x44)

        out_path = reports_dir / f"{bare}_{timestamp}.docx"
        doc.save(str(out_path))
    finally:
        # 清掉暫存圖（保留 docx 即可）
        for p in tmp_dir.glob("*.png"):
            try:
                p.unlink()
            except Exception:
                pass
        try:
            tmp_dir.rmdir()
        except Exception:
            pass

    return out_path


def main() -> int:
    import argparse
    parser = argparse.ArgumentParser(
        description="股票溝通師 DOCX 報告產生器（V3.3）",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""\
範例（指令一律一行寫完，不要用 bash `\\` 或 PowerShell 反引號做行繼續）：
  # 基本用法（自動生成擬人化和總結）
  python generate_report.py 2330 "台積電,台積,TSMC"

  # 由 Claude 在對話中提供擬人化和總結（最有溫度）——
  # 先把文字寫進檔案再用 --persona-file / --summary-file 讀，跨平台最穩
  # （路徑可自行決定，建議放 skill 目錄下的 tmp/）
  python generate_report.py 2330 "台積電,台積" --persona-file ./tmp/persona.txt --summary-file ./tmp/summary.txt
""",
    )
    parser.add_argument("symbol", help="股票代號（例：2330、AAPL、00631L）")
    parser.add_argument("name_zh", nargs="?", default=None,
                        help="中文別名清單（逗號分隔）")
    parser.add_argument("--persona", default=None,
                        help="擬人化獨白（Claude 對話中產生）。沒提供就用模板自動生成")
    parser.add_argument("--summary", default=None,
                        help="一句話總結（Claude 對話中產生）。沒提供就用模板自動生成")
    parser.add_argument("--persona-file", default=None,
                        help="從檔案讀取擬人化獨白（適合長文字）")
    parser.add_argument("--summary-file", default=None,
                        help="從檔案讀取一句話總結")
    parser.add_argument("--output-dir", default=None,
                        help="報告輸出目錄。預設為當前工作目錄下的 reports/。"
                             "建議由 Claude 帶入使用者專案目錄，避免污染共享 skill 資料夾。")
    args = parser.parse_args()

    # 從檔案讀（如果有指定）
    persona = args.persona
    if args.persona_file:
        from pathlib import Path as _Path
        persona = _Path(args.persona_file).read_text(encoding="utf-8")

    summary = args.summary
    if args.summary_file:
        from pathlib import Path as _Path
        summary = _Path(args.summary_file).read_text(encoding="utf-8")

    out = build_report(args.symbol, args.name_zh, persona=persona, summary=summary,
                        output_dir=args.output_dir)
    print(f"\n✅ 報告產出：{out}")
    print(f"   檔案大小：{out.stat().st_size / 1024:.1f} KB")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
